import os
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google.generativeai.types import tool
class ResumeOptimizer:
    def __init__(self, llm):
        self.llm = llm
    @tool
    def generate(self, target_role: str, candidate_name: str, resume_text: str = None):
        """
        Optimize the resume for a given target role and save it to Desktop as:
        optimized_resume_<role>_<name>.docx
        """
        if not resume_text:
            return "⚠️ No resume text found. Please load a resume first."

        # ✅ Create clean role + name for filename
        safe_role = "_".join(target_role.split())
        safe_name = "".join(c for c in candidate_name if c.isalnum())

        # ✅ File path on Desktop
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        os.makedirs(desktop, exist_ok=True)
        filename = f"optimized_resume_{safe_role}_{safe_name}.docx"
        out_path = os.path.join(desktop, filename)

        # ✅ Prompt to LLM
        prompt = f"""
        You are Resumini, an expert AI resume optimization agent trained in HR standards, job market keywords, and ATS scoring systems.
        Your goal is to rewrite the given resume to maximize ATS compatibility and recruiter appeal for the role of **{target_role}**.

        Rules:
        - Use only text-based formatting (no tables, columns, or graphics).
        - Keep consistent professional tone and layout (ATS-friendly headings like SUMMARY, SKILLS, PROJECTS, EXPERIENCE, EDUCATION).
        - Emphasize keywords relevant to {target_role}.
        - Use strong action verbs and concise bullet points.
        - Quantify impact where possible.
        - Avoid repetition or filler text.
        - Maintain factual accuracy.
        - Output ONLY the final resume text — do not include commentary or meta notes.

        Input Resume:
        --------------------
        {resume_text}

        Output:
        --------------------
        ATS-Optimized Resume for {target_role}:
        """

        try:
            optimized_text = self.llm.generate(prompt)
        except Exception as e:
            return f"⚠️ LLM generation failed: {e}"

        if not optimized_text or "Error" in optimized_text:
            return "⚠️ Optimization failed. Please check your API key or LLM response."

        # ✅ Create and format Word document
        try:
            doc = docx.Document()
            title = doc.add_heading(f"Optimized Resume for {target_role}", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for line in optimized_text.splitlines():
                line = line.strip()
                if not line:
                    continue
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

            doc.save(out_path)
        except Exception as e:
            return f"⚠️ Error saving resume: {e}"

        print(f"✅ Resume optimized and saved to Desktop as: {filename}")
        return out_path
